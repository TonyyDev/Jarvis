"""Funciones para trabajar con documentos Word."""

from openpyxl import Workbook
import openpyxl
from docx import Document
import webbrowser


def crear_word(nombre_archivo):
    documento = Document()

    documento.add_heading(
        "Documento creado por Jarvis",
        level=1
    )

    documento.add_paragraph(
        "Este documento fue creado automáticamente utilizando Python."
    )

    documento.save(nombre_archivo)

    print("Documento creado correctamente.")


def leer_word(nombre_archivo):
    documento = Document(nombre_archivo)

    print("Contenido del documento:")
    print()

    for parrafo in documento.paragraphs:
        print(parrafo.text)


def agregar_texto_word(nombre_archivo, texto):
    documento = Document(nombre_archivo)

    documento.add_paragraph(texto)

    documento.save(nombre_archivo)

    print("Texto agregado correctamente.")




# crear_word("informe.docx")

# agregar_texto_word(
#     "informe.docx",
#     "Este párrafo fue agregado posteriormente como un tipo informe por Jarvis."
# )

# leer_word("informe.docx")






def crear_excel(nombre_archivo):
    libro = Workbook()

    hoja = libro.active
    hoja.title = "Ventas"

    hoja["A1"] = "Producto"
    hoja["B1"] = "Cantidad"
    hoja["C1"] = "Precio"

    hoja["A2"] = "Pan"
    hoja["B2"] = 10
    hoja["C2"] = 500

    hoja["A3"] = "Pastel"
    hoja["B3"] = 5
    hoja["C3"] = 3500

    libro.save(nombre_archivo)

    print("Excel creado correctamente.")




    
#crear_excel("ventas.xlsx")



def leer_excel(nombre_archivo):
    libro = openpyxl.load_workbook(nombre_archivo)

    hoja = libro.active

    print("Contenido del Excel:")
    print()

    for fila in hoja.iter_rows(values_only=True):
        print(fila)



# leer_excel("ventas.xlsx")

def abrir_youtube():
    webbrowser.open("https://www.youtube.com")
    print("YouTube abierto correctamente.")

# abrir_youtube()



def abrir_navegador():
    webbrowser.open("https://www.google.com")
    print("google abierto correctamente.")

abrir_navegador()